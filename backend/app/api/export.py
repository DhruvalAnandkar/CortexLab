"""
Export API Routes

Word document export functionality.
"""

from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

from app.core.database import get_db
from app.dependencies import get_current_user
from app.models import User, Project, Artifact

router = APIRouter()


def markdown_to_docx(markdown_content: str, title: str) -> BytesIO:
    """
    Convert markdown content to a Word document.
    
    Handles basic markdown formatting:
    - Headers (#, ##, ###)
    - Bold (**text**)
    - Italic (*text*)
    - Lists (- or *)
    - Paragraphs
    """
    doc = Document()
    
    # Set document title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Split content into lines
    lines = markdown_content.split("\n")
    current_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:
            if current_list:
                current_list = False
            continue
        
        # Handle headers
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        # Handle lists
        elif line.startswith("- ") or line.startswith("* "):
            text = line[2:]
            text = format_inline_text(text)
            p = doc.add_paragraph(text, style="List Bullet")
            current_list = True
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            text = format_inline_text(text)
            p = doc.add_paragraph(text, style="List Number")
            current_list = True
        else:
            # Regular paragraph
            text = format_inline_text(line)
            p = doc.add_paragraph(text)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def format_inline_text(text: str) -> str:
    """
    Format inline markdown (bold, italic) for docx.
    
    Note: This is a simplified version. For production, 
    you'd want to use runs with proper formatting.
    """
    # Remove markdown formatting for now (simplified)
    # In production, you'd parse and apply Word formatting
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Remove bold markers
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # Remove italic markers
    text = re.sub(r"`(.+?)`", r"\1", text)  # Remove code markers
    return text


@router.post("/artifacts/{artifact_id}/export/docx")
async def export_to_docx(
    artifact_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Export an artifact to a Word document (.docx).
    
    Converts the markdown content to a formatted Word document.
    """
    # Get artifact
    result = await db.execute(
        select(Artifact).where(Artifact.id == artifact_id)
    )
    artifact = result.scalar_one_or_none()
    
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")
    
    # Verify ownership through project
    project_result = await db.execute(
        select(Project)
        .where(Project.id == artifact.project_id, Project.user_id == current_user.id)
    )
    if not project_result.scalar_one_or_none():
        raise HTTPException(status_code=404, detail="Artifact not found")
    
    # Convert to docx
    docx_buffer = markdown_to_docx(artifact.content_markdown, artifact.title)
    
    # Generate safe filename
    safe_title = re.sub(r"[^\w\s-]", "", artifact.title)[:50]
    filename = f"{safe_title}.docx"
    
    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )
